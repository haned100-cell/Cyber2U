from docx import Document

DOCX_PATH = '/Users/nathanbrown-bennett/Haned-Cyber2U/Final Disso/Cyber2U_finished_dissertation_final.docx'
START = 104
END = 210


def main() -> None:
    doc = Document(DOCX_PATH)
    for idx in range(START, min(END, len(doc.paragraphs))):
        p = doc.paragraphs[idx]
        text = p.text.replace('\n', ' ').strip()
        style = p.style.name if p.style is not None else ''
        print(f'{idx:04d} | {style} | {text}')


if __name__ == '__main__':
    main()
