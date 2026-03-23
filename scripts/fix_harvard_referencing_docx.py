from pathlib import Path
from docx import Document

DOCX_PATH = Path('/Users/nathanbrown-bennett/Haned-Cyber2U/FILE_0215.docx')

if not DOCX_PATH.exists():
    raise FileNotFoundError(f'Document not found: {DOCX_PATH}')

SOURCE_LINE = "Source: Author's own screenshots (2026)."

# Replace outdated in-text figure references with updated figure mapping.
TEXT_REPLACEMENTS = {
    "open http://localhost:8025          # MailHog email UI (see Figures 4 and 5 in Appendix A)": "open http://localhost:8025          # MailHog email UI (see Figures 6 and 7 in Appendix A)",
    "QuizPlayer.tsx — Interactive quiz interface (see Figure 3 in Appendix A)": "QuizPlayer.tsx — Interactive quiz interface and review flow (see Figures 5, 8 and 9 in Appendix A)",
    "LearnerDashboard.tsx — Progress tracking and insights (see Figure 2 in Appendix A)": "LearnerDashboard.tsx — Progress tracking and insights (see Figures 2, 3 and 4 in Appendix A)",
}

REFERENCES_HEADING = "References"
REFERENCE_ITEMS = [
    "Bandura, A. (1977) Social Learning Theory. Englewood Cliffs, NJ: Prentice Hall.",
    "Bonneau, J., Herley, C., Van Oorschot, P.C. and Stajano, F. (2012) 'The quest to replace passwords: A framework for comparative evaluation of web authentication schemes', 2012 IEEE Symposium on Security and Privacy, pp. 553-567. doi: 10.1109/SP.2012.44.",
    "Cofense (2023) Cofense phishing awareness and reporting solutions. Available at: https://cofense.com/ (Accessed: 23 March 2026).",
    "Proofpoint (2023) Security awareness training platform. Available at: https://www.proofpoint.com/ (Accessed: 23 March 2026).",
]


def replace_text_if_exact(paragraph_text: str) -> str:
    return TEXT_REPLACEMENTS.get(paragraph_text, paragraph_text)


def set_paragraph_text(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return

    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def has_source_after(paragraphs, idx: int) -> bool:
    if idx + 1 >= len(paragraphs):
        return False
    next_text = paragraphs[idx + 1].text.strip()
    return next_text.startswith("Source:")


def add_sources_below_figures(document: Document) -> None:
    paragraphs = document.paragraphs
    insert_positions = []

    for idx, p in enumerate(paragraphs):
        text = p.text.strip()
        if text.startswith("Figure ") and "." in text:
            if not has_source_after(paragraphs, idx):
                insert_positions.append(idx)

    # Insert from bottom to top to preserve indices.
    for idx in reversed(insert_positions):
        if idx + 1 < len(document.paragraphs):
            source_para = document.paragraphs[idx + 1].insert_paragraph_before(SOURCE_LINE)
        else:
            source_para = document.add_paragraph(SOURCE_LINE)

        try:
            source_para.style = document.paragraphs[idx].style
        except Exception:
            pass


def ensure_references_section(document: Document) -> None:
    paragraphs = document.paragraphs
    existing_headings = {p.text.strip().lower(): p for p in paragraphs}

    if REFERENCES_HEADING.lower() not in existing_headings:
        document.add_paragraph("")
        heading = document.add_paragraph(REFERENCES_HEADING)
        try:
            heading.style = 'Heading 1'
        except Exception:
            pass

    existing_texts = {p.text.strip() for p in document.paragraphs}
    for ref in REFERENCE_ITEMS:
        if ref not in existing_texts:
            document.add_paragraph(ref)


def main() -> None:
    doc = Document(str(DOCX_PATH))

    for p in doc.paragraphs:
        old_text = p.text.strip()
        new_text = replace_text_if_exact(old_text)
        if new_text != old_text:
            set_paragraph_text(p, new_text)

    add_sources_below_figures(doc)
    ensure_references_section(doc)

    doc.save(str(DOCX_PATH))
    print('Updated Harvard-style references, in-text figure links, and figure sources in FILE_0215.docx')


if __name__ == '__main__':
    main()
